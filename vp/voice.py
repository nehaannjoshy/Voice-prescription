import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import speech_recognition as sr
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm , RGBColor , Inches
from docx.shared import Pt 
from tkinter import *
import smtplib 
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText 
from email.mime.base import MIMEBase 
from email import encoders
from datetime import date
from stat import S_IREAD, S_IRGRP, S_IROTH
from tkinter import ttk
import exception

r = tk.Tk() 

def verify(name=None):
	if(name!=None):
		p="C:/Users/JOEL/Desktop/apptest/vp/{}.docx".format(name)
	else:
		r.filename =  filedialog.askopenfilename(initialdir = "C:/Users/JOEL/Desktop/apptest/vp",title = "Select file",filetypes = (("document files","*.docx"),("all files","*.*")))
		p=r.filename
	a=os.path.split(p)
	os.system(a[1])
	#os.chmod(a[1], S_IREAD|S_IRGRP|S_IROTH)
	upload(p)
	


def doc():
	ri=tk.Tk() 
	ri.filename =  filedialog.askopenfilename(initialdir = "C:/Users/JOEL/Desktop/apptest/vp",title = "Select file",filetypes = (("wav files","*.wav"),("all files","*.*")))
	
	a=os.path.split(ri.filename)
	
	r=sr.Recognizer()
	with sr.AudioFile(a[1]) as source:
		audio=r.listen(source)
	#try:	
		text=r.recognize_google(audio)
		messagebox.showinfo("Alert","Converted Text: {}".format(text))
		text=text.lower()
		str=text.split(" ")
		val="mg"
		for i in range(str.count(val)):
			str.remove(val)
		val="milligram"
		for i in range(str.count(val)):
			str.remove(val)		
		val="days"
		for i in range(str.count(val)):
			str.remove(val)
		val="ml"
		for i in range(str.count(val)):
			str.remove(val)
		val="for"
		for i in range(str.count(val)):
			str.remove(val)

		name=""
		age=""
		Date=date.today()
		ad=""
		dia=""
		sym=""
		tablet=[]
		syrup=[]
		for i in range(len(str)):
		    if(str[i]=='name'):
		        name=str[i+2]
		    if(str[i]=='i' and str[i+1]=='am'):
		    	name=str[i+2]
		    if(str[i]=="i'm"):
		    	name=str[i+2]
		    if(str[i]=='name' and str[i+3]!='age'):
		    	name+=str[i+3]
		    if(str[i]=='age'):
		        age=str[i+1]
		    if(str[i]=='tablet'):
		        ar=[]
		        default="0-0-0"
		        ar.append(str[i+1])
		        if(str[i+2].find("mg")>=0):
		        	str[i+2]=str[i+2][:-2]

		        ar.append(str[i+2])
				

		        if(str[i+3].find("days")>=0):
		        	str[i+3]=str[i+3][:-4]


		        ar.append(str[i+3])


		        if(str[i+4]=='before'):
		            ar.append(str[i+4])	     	
		        else:
		            ar.append("After")
		        for j in range(5):
		        	if(str[i+j+3]=='once'):
		        		default='1'+'-'+default[2]+'-'+default[4]
		        		break
		        	if(str[i+j+3]=='twice'):
		        		default='1'+'-'+default[2]+'-'+'1'
		        		break
		        	if(str[i+j+3]=='thrice' or str[i+j+3]=='prices' or str[i+j+3]=='rice'):
		        		default='1'+'-'+'1'+'-'+'1'
		        		break
		        	if(str[i+j+3]=='night'):
		        		default=default[0]+'-'+default[2]+'-'+'1'
		        	if(str[i+j+3]=='afternoon'):
		        		default=default[0]+'-'+'1'+'-'+default[4]
		        	if(str[i+j+3]=='morning'):
		        		default='1'+'-'+default[2]+'-'+default[4]
		        	if(str[i+j+3]=='done' or str[i+j+3]=='dun'):
		        		break

		            
		        
		        ar.append(default)
		        tablet.append(ar)
		    if(str[i]=='syrup'):
		        arr=[]
		        default="0-0-0"
		        arr.append(str[i+1])


		        if(str[i+2].find("ml")>=0):
		        	str[i+2]=str[i+2][:-2]
		        	
		        
		        arr.append(str[i+2])
		        
		        if(str[i+3].find("days")>=0):
		        	str[i+3]=str[i+3][:-4]
		        	
		        
		        arr.append(str[i+3])

		        


		        if(str[i+4]=='before'):
		            arr.append(str[i+4])
		        else:
		            arr.append("After")
		        for j in range(8):
		            if(str[i+j+3]=='once'):
		            	default='1'+'-'+default[2]+'-'+default[4]
		            	break
		            if(str[i+j+3]=='twice'):
		            	default='1'+'-'+default[2]+'-'+'1'
		            	break
		            if(str[i+j+3]=='thrice' or str[i+j+3]=='prices'):
		            	default='1'+'-'+'1'+'-'+'1'
		            	break
		            if(str[i+j+3]=='night'):
		            	default=default[0]+'-'+default[2]+'-'+'1'
		            if(str[i+j+3]=='afternoon'):
		            	default=default[0]+'-'+'1'+'-'+default[4]
		                
		            if(str[i+j+3]=='morning'):
		                
		                default='1'+'-'+default[2]+'-'+default[4]
		            if(str[i+j+3]=='done' or str[i+j+3]=='dun'):
		            	break



		            
		        
		        arr.append(default)
		        syrup.append(arr)
		        


		if((len(name))==0):
			messagebox.showinfo("Error","No patient name is found, please retry again!")
			raise Exception('No name found')	            

		document=Document()
		para=document.add_heading("")
		runn=para.add_run("ABC HOSPITAL \n")
		run=para.add_run(" Dr.Jane Doe M.D.(Neurology) \t \t Dr.Smith M.S.(Ortho) \t Dr.Roshni M.S.(Ophthalmology) \n Sector: 17 Navi Mumbai \n Ph:0987654321 \n" )
		font = run.font
		runn.font.size=Pt(18)
		runn.font.color.rgb=RGBColor(0,0,0)
		font.size = Pt(16)
		font.color.rgb = RGBColor(0,0,0)
		para.alignment = WD_ALIGN_PARAGRAPH.CENTER

		para1=document.add_paragraph()
		run1=para1.add_run("Name:{}".format(name))
		run1.font.size=Pt(14)


		para2=document.add_paragraph()
		run2=para2.add_run("Age:{}".format(age))
		run2.font.size=Pt(14)

		para3=document.add_paragraph()
		run3=para3.add_run("Date:{}".format(Date))
		run3.font.size=Pt(14)



		if(len(tablet)==0):
			messagebox.showinfo("Alert","Please include atleast one tablet")
		else:
			table = document.add_table(len(tablet), 5)

			heading_cells = table.rows[0].cells
			heading_cells[0].text = 'Tablet name'
			heading_cells[1].text = 'dosage(mg)'
			heading_cells[2].text = 'No.of.days'
			heading_cells[3].text = 'Before or after food'
			heading_cells[4].text = 'No.of.times per day'

			for i in range(len(tablet)):
			    cells = table.add_row().cells
			    cells[0].text = tablet[i][0]
			    cells[1].text = tablet[i][1]
			    cells[2].text = tablet[i][2]
			    cells[3].text = tablet[i][3]
			    cells[4].text = tablet[i][4]
		#for space
		para4=document.add_paragraph("\n \n")
		if(len(syrup)!=0):
			table1 = document.add_table(len(syrup), 5)
			heading_cells = table1.rows[0].cells
			heading_cells[0].text = 'Syrup name'
			heading_cells[1].text = 'dosage(ml)'
			heading_cells[2].text = 'No.of.days'
			heading_cells[3].text = 'Before or after food'
			heading_cells[4].text = 'No.of.times per day'
			    # add a data row for each item
			for i in range(len(syrup)):
			    cells = table1.add_row().cells
			    cells[0].text = syrup[i][0]
			    cells[1].text = syrup[i][1]
			    cells[2].text = syrup[i][2]
			    cells[3].text = syrup[i][3]
			    cells[4].text = syrup[i][4]                               
                              
                                        

		para5=document.add_paragraph("\n \n")



		
		header = document.sections[0].footer
		f1=header.add_paragraph()
		
		runf=f1.add_run("\n Signature \n I hereby accept that this prescription was verified")
		runf.font.size=Pt(16)
		runf.font.color.rgb=RGBColor(0,0,0)
		f1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		
		document.save("{}.docx".format(name))
		messagebox.showinfo("Successfully completed", "Success! Please verify the document generated and click upload \n Document Name:{}.docx".format(name))
		

		verify(name)
	#except:
		#messagebox.showinfo("Error","please try again!")
	
		ri.destroy()
	



def listen():
	r = sr.Recognizer()
	with sr.Microphone() as source:
		messagebox.showinfo("Alert","Start Speaking")
		audio = r.listen(source)
		messagebox.showinfo("Alert","completed.")

	text=r.recognize_google(audio)
	messagebox.showinfo("Alert","Converted Text:{}".format(text))
	text=text.lower()
	str=text.split(" ")
	val="mg"
	var="milligram"
	for i in range(str.count(val)):
		str.remove(val)
	for i in range(str.count(val)):
		str.remove(var)
	val="days"
	for i in range(str.count(val)):
		str.remove(val)
	val="ml"
	for i in range(str.count(val)):
		str.remove(val)
	val="for"
	for i in range(str.count(val)):
		str.remove(val)
	name=""
	age=""
	Date=date.today()
	ad=""
	dia=""
	sym=""
	tablet=[]
	syrup=[]
	for i in range(len(str)):
	    if(str[i]=='name'):
	        name=str[i+2]

	    if(str[i]=='name' and str[i+3]!='age'):
	    	name+=str[i+3]
	    if(str[i]=='age'):
	        age=str[i+1]
	    if(str[i]=='tablet'):
	        ar=[]
	        default="0-0-0"
	        ar.append(str[i+1])
	        if(str[i+2].find("mg")>=0):
	        	str[i+2]=str[i+2][:-2]
	        ar.append(str[i+2])
			
	        if(str[i+3].find("days")>=0):
	        	str[i+3]=str[i+3][:-4]
	        ar.append(str[i+3])
	        if(str[i+4]=='before'):
	            ar.append(str[i+4])	     	
	        else:
	            ar.append("After")
	        for j in range(5):
	            if(str[i+j+3]=='once'):
	            	default='1'+'-'+default[2]+'-'+default[4]
	            	break
	            if(str[i+j+3]=='twice'):
	            	default='1'+'-'+default[2]+'-'+'1'
	            	break
	            if(str[i+j+3]=='thrice'):
	            	default='1'+'-'+'1'+'-'+'1'
	            	break
	            if(str[i+j+3]=='night'):
	            	default=default[0]+'-'+default[2]+'-'+'1'
	            if(str[i+j+3]=='afternoon'):
	            	default=default[0]+'-'+'1'+'-'+default[4]
	                
	            if(str[i+j+3]=='morning'):
	                
	                default='1'+'-'+default[2]+'-'+default[4]
	            if(str[i+j+3]=='done' or str[i+j+3]=='dun'):
	            	break	        	
	        ar.append(default)
	        tablet.append(ar)
	    if(str[i]=='syrup'):
	        arr=[]
	        default="0-0-0"
	        arr.append(str[i+1])
	        if(str[i+2].find("ml")>=0):
	        	str[i+2]=str[i+2][:-2]
	        	
	        
	        arr.append(str[i+2])
	        
	        if(str[i+3].find("days")>=0):
	        	str[i+3]=str[i+3][:-4]
	        	
	        
	        arr.append(str[i+3])
	        
	        if(str[i+4]=='before'):
	            arr.append(str[i+4])
	        else:
	            arr.append("After")
	        for j in range(8):
	            if(str[i+j+3]=='once'):
	            	default='1'+'-'+default[2]+'-'+default[4]
	            	break
	            if(str[i+j+3]=='twice'):
	            	default='1'+'-'+default[2]+'-'+'1'
	            	break
	            if(str[i+j+3]=='thrice'):
	            	default='1'+'-'+'1'+'-'+'1'
	            	break
	            if(str[i+j+3]=='night'):
	            	default=default[0]+'-'+default[2]+'-'+'1'
	            if(str[i+j+3]=='afternoon'):
	            	default=default[0]+'-'+'1'+'-'+default[4]
	                
	            if(str[i+j+3]=='morning'):
	                
	                default='1'+'-'+default[2]+'-'+default[4]
	            if(str[i+j+3]=='done' or str[i+j+3]=='dun'):
	            	break
	            
	        
	        arr.append(default)
	        syrup.append(arr)  

	if((len(name))==0):
		messagebox.showinfo("Error","No patient name is found, please retry again!")
		raise Exception('No name found')
	document=Document()
	para=document.add_heading("")
	runn=para.add_run("ABC HOSPITAL \n")
	run=para.add_run(" Dr.Jane Doe M.D.(Neurology) \t \t Dr.Smith M.S.(Ortho) \t Dr.Roshni M.S.(Ophthalmology) \n Sector: 17 Navi Mumbai \n Ph:0987654321 \n" )
	font = run.font
	runn.font.size=Pt(18)
	runn.font.color.rgb=RGBColor(153, 17, 150)
	font.size = Pt(16)
	font.color.rgb = RGBColor(217, 17, 213)
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	para1=document.add_paragraph()
	run1=para1.add_run("Name:{}".format(name))
	run1.font.size=Pt(14)
	para2=document.add_paragraph()
	run2=para2.add_run("Age:{}".format(age))
	run2.font.size=Pt(14)
	para3=document.add_paragraph()
	run3=para3.add_run("Date:{}".format(Date))
	run3.font.size=Pt(14)


	if(len(tablet)==0):
	    messagebox.showinfo("Alert","Please include atleast one tablet")
	    raise Exception('No name found')
	else:
		table = document.add_table(len(tablet), 5)
	
		heading_cells = table.rows[0].cells
		heading_cells[0].text = 'Tablet name'
		heading_cells[1].text = 'dosage(mg)'
		heading_cells[2].text = 'No.of.days'
		heading_cells[3].text = 'Before or after food'
		heading_cells[4].text = 'No.of.times per day'
		    # add a data row for each item
		for i in range(len(tablet)):
		    cells = table.add_row().cells
		    cells[0].text = tablet[i][0]
		    cells[1].text = tablet[i][1]
		    cells[2].text = tablet[i][2]
		    cells[3].text = tablet[i][3]
		    cells[4].text = tablet[i][4]
	#for space
	para4=document.add_paragraph("\n \n")
	if(len(syrup)!=0):
	    #syrup
	    table1 = document.add_table(len(syrup), 5)
	        # populate header row --------
	    heading_cells = table1.rows[0].cells
	    heading_cells[0].text = 'Syrup name'
	    heading_cells[1].text = 'dosage(ml)'
	    heading_cells[2].text = 'No.of.days'
	    heading_cells[3].text = 'Before or after food'
	    heading_cells[4].text = 'No.of.times per day'
	        # add a data row for each item
	    for i in range(len(syrup)):
	        cells = table1.add_row().cells
	        cells[0].text = syrup[i][0]
	        cells[1].text = syrup[i][1]
	        cells[2].text = syrup[i][2]
	        cells[3].text = syrup[i][3]
	        cells[4].text = syrup[i][4]                               
	                      
	                                
	para5=document.add_paragraph("\n \n")

	header = document.sections[0].footer
	f1=header.add_paragraph()
	
	runf=f1.add_run("\n Signature \n I hereby accept that this prescription was verified")
	runf.font.size=Pt(16)
	runf.font.color.rgb=RGBColor(0,0,0)
	f1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	document.save("{}.docx".format(name))

	
def upload(p=None):
	u=tk.Tk()
	u.geometry('400x400')
	u.configure(background = '#54a8cc');
	L1 = Label(u, text="Patient E-mail",bg="#FFFFBF")
	L1.pack( side = LEFT)
	E1 = Entry(u, bd =5, width="30")
	E1.pack(side = RIGHT)
	'''L2 = Label(u, text="Password",bg="#ef47f5")
	L2.pack( side = LEFT)
	E2 = Entry(u, bd =5, width="30")
	E2.pack(side = RIGHT)
	L3 = Label(u, text="Patient email",bg="#ef47f5")
	L3.pack( side = LEFT)
	E3 = Entry(u, bd =5, width="30")
	E3.pack(side = RIGHT)'''
	
	def upload1(p):
		toaddr=E1.get()		
		fromaddr='#emailid' #Enter Your E-mail id 
		msg = MIMEMultipart() 
		msg['From'] = fromaddr 
		msg['To'] = toaddr 
		msg['Subject'] = "Voice prescription"
		body = "Prescription"
		msg.attach(MIMEText(body, 'plain')) 
		if(p==None):
			u.filename =  filedialog.askopenfilename(initialdir = "C:/Users/JOEL/Desktop/apptest/p",title = "Select file",filetypes = (("document files","*.docx"),("all files","*.*")))
	
			a=os.path.split(u.filename)
			print(a)
			filename = a[1]
			addr=a[0]+'/{}'.format(a[1])
		else:
			a=os.path.split(p)
			print(a)
			filename = a[1]
			addr=a[0]+'/{}'.format(a[1])
		attachment = open(addr, "rb") 
		p = MIMEBase('application', 'octet-stream')
		p.set_payload((attachment).read()) 
		encoders.encode_base64(p) 
		p.add_header('Content-Disposition', "attachment; filename= %s" % filename) 
		msg.attach(p) 
		s = smtplib.SMTP('smtp.gmail.com', 587) 
		s.starttls() 
		s.login(fromaddr,'#password' )  #Enter Your email password
		text = msg.as_string() 
		s.sendmail(fromaddr, toaddr, text) 
		top=tk.Tk()
		messagebox.showinfo("Successfully completed", "Prescription sent successfully")
		top.mainloop()
		top.destroy()
		s.quit() 
		u.destroy()


	submit = Button(u, text ="Submit", command =lambda: upload1(p),bg="#FFFFBF")
	submit.pack(side =BOTTOM) 



r.geometry('500x500')
var = StringVar()
label = Label( r,bg='#54a8cc',textvariable=var, relief=RAISED )
label.config(font=("Courier", 44))

var.set("Dr. Py")
label.pack()
r.configure(background = '#54a8cc');
r.title('Online Prescription')
try:
        btn1 = tk.Button(r, text='Upload conversation', width=25, command=doc,bg="#ffffff") 

        btn2 = tk.Button(r, text='Start new recording', width=25, command=listen,bg="#ffffff") 
        btn4=tk.Button(r, text='Verify Prescription', width=25, command=verify,bg="#ffffff") 
        btn3 = tk.Button(r, text='Send Prescription through mail', width=25, command=upload,bg="#ffffff")
except:
        messagebox.showinfo("Error", "Something went wrong, please try again")
btn1.pack(pady=20) 
btn2.pack(pady=20)
btn4.pack(pady=20)
btn3.pack(pady=20) 

r.mainloop()
